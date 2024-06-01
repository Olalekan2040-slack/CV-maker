from flask import Flask, render_template, request, send_file
from docx import Document
from docx.shared import Pt, RGBColor
from io import BytesIO
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
app = Flask(__name__)

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        full_name = request.form['full_name']
        location = request.form['location']
        email = request.form['email']
        title = request.form['title']
        template = request.form['template']


        doc = create_doc(full_name,location, email, title, template)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return send_file(buffer, as_attachment=True, download_name=full_name + ".docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    
    return render_template('index.html')


@app.route('/preview', methods=['POST'])
def preview():
    full_name = request.form['full_name']
    location = request.form['location']
    email = request.form['email']
    title = request.form['title']
    template = request.form['template']


    doc = create_doc(full_name,location,email,title, template)

    return render_template('preview.html', full_name=full_name,location=location, email=email, title=title, template=template)


def create_doc(full_name,location,email,title, template):
    doc = Document()

    if template == "template1":
        name = doc.add_heading(full_name, level=1)
        name.alignment = 1

        run = name.runs[0]
        run.font.name = 'Calibri'
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

        #set the background colour to light Yellow
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:val'), 'clear')
        shading_elm.set(qn('w:color'), 'auto')
        shading_elm.set(qn('w:fill'), 'FFFF99')
        run._element.get_or_add_rPr().append(shading_elm)

        #set location, title alignment
        p = doc.add_paragraph(location, style="Normal")
        p.alignment = 1
        p.add_run(" | ")
        p.add_run(email).italic=True
        p.add_run(" | ")
        p.add_run(title)
        p_format = p.paragraph_format
        p_format = Pt(12)



        #Add a horinzontal line after heading
        p_single_line = doc.add_paragraph()
        p_single_line.alignment = 1
        p_format = p_single_line.paragraph_format
        p_format = Pt(12)
        p_border = p_single_line._element.get_or_add_pPr()

        # Define the border properties for a single bottom line
        border_attrs = {
            'w:val': 'single',
            'w:sz': '4',  # Size of the line
            'w:space': '1',  # Space between text and line
            'w:color': '000000'  # Color of the line (black)
        }

        # Add the bottom border
        border_tag = 'w:bottom'
        border_elm = OxmlElement(border_tag)
        for attr, value in border_attrs.items():
            border_elm.set(qn(attr), value)
        p_border.append(border_elm)







    elif template == 'template2':
        pass

    elif template == "template3":
        pass
    
    return doc





if __name__ == '__main__':
    app.run(debug=True)